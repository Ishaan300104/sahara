from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_field(doc, label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_val = p.add_run(value)
    run_val.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

def add_rating_question(doc, qnum, qtext, rating, rating_labels):
    p = doc.add_paragraph()
    r = p.add_run(f"{qnum}. {qtext}")
    r.bold = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)

    p2 = doc.add_paragraph()
    for i, label in enumerate(rating_labels, 1):
        box = "[X]" if i == rating else "[ ]"
        p2.add_run(f"  {box} {i} - {label}   ")
    p2.runs[0].font.size = Pt(9)
    p2.paragraph_format.space_after = Pt(4)

def add_text_question(doc, qnum, qtext, answer):
    p = doc.add_paragraph()
    r = p.add_run(f"{qnum}. {qtext}")
    r.bold = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)

    p2 = doc.add_paragraph(answer)
    p2.runs[0].font.size = Pt(10)
    p2.paragraph_format.space_after = Pt(6)

def add_divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    p.paragraph_format.space_after = Pt(4)

respondents = [
    {
        "name": "Ishaan Pandey",
        "date": "18th April 2026",
        "familiarity": "Intermediate",
        "duration": "15 minutes",
        "q1": 5, "q2": 4, "q5": 5, "q6": 4, "q8": 3, "q11": 4, "q16": 4, "q22": 3, "q28": 4,
        "q3": "The mood history feature was slightly hard to locate at first. Navigation to past sessions could be more intuitive.",
        "q4": "A journaling feature and weekly mood summary would be helpful.",
        "q7": "Once the chatbot responded with a slightly generic coping tip that felt out of context, but overall it was fine.",
        "q9": "Yes, I could end the session or rephrase my input. No undo for AI responses though.",
        "q10": "Only text inputs during conversation — no name, location, or health records required.",
        "q12": "A brief disclaimer was shown at the start stating data is not stored beyond the session. Clear enough.",
        "q13": "No unnecessary permissions were requested.",
        "q14": "No biased outputs observed.",
        "q15": "It should work reasonably well across backgrounds, but responses felt tuned to English-speaking, urban users.",
        "q17": "Not always clear. The chatbot blends AI and scripted responses without distinguishing them.",
        "q18": "No explicit error-reporting mechanism. You can just rephrase and continue.",
        "q19": "It states at the start it is not a therapist, which is important.",
        "q20": "A pop-up consent message appeared before the first chat session.",
        "q21": "Yes — you can simply not engage with the mood tracking prompts.",
        "q23": "If a user in genuine crisis relies solely on this app instead of seeking real help, that is a serious risk.",
        "q24": "No harmful outputs encountered.",
        "q25": "Users with severe mental illness could be harmed if they over-rely on this app.",
        "q26": "Could reduce referrals to student counselors if students feel the chatbot is sufficient.",
        "q27": "Minimal for a prototype, but LLM inference at scale has energy implications.",
        "q29": "Add a clearer crisis escalation path with real counselor contact.",
        "q30": "Good effort for a course project. The empathetic tone is well done.",
    },
    {
        "name": "Deekshant Singh Rajawat",
        "date": "18th April 2026",
        "familiarity": "Intermediate",
        "duration": "20 minutes",
        "q1": 4, "q2": 4, "q5": 4, "q6": 4, "q8": 2, "q11": 4, "q16": 3, "q22": 3, "q28": 4,
        "q3": "Nothing was broken, but the response loading time was occasionally slow.",
        "q4": "Voice input option and multi-language support would broaden accessibility.",
        "q7": "One response seemed overly optimistic for a message expressing hopelessness — a bit tone-deaf.",
        "q9": "Limited control — I could not modify or delete a previous message in the conversation.",
        "q10": "Only chat messages. No personal identifiers collected.",
        "q12": "Brief disclaimer at the top of the chat. No full privacy policy linked.",
        "q13": "No excessive data requests noted.",
        "q14": "No obvious bias, but responses felt slightly Western in framing of mental health concepts.",
        "q15": "May not resonate equally with users from different cultural contexts around mental health.",
        "q17": "It is not transparent about when it is following a script vs. generating a dynamic response.",
        "q18": "No reporting mechanism exists currently.",
        "q19": "The disclaimer is there but could be more prominent throughout the conversation.",
        "q20": "A checkbox consent before starting the session.",
        "q21": "No explicit opt-out button, but you can ignore AI suggestions.",
        "q23": "The app could give false reassurance to someone who needs urgent clinical help.",
        "q24": "No. Responses stayed within appropriate boundaries.",
        "q25": "People with active mental health conditions and those with low digital literacy are more at risk.",
        "q26": "Could partially substitute peer support roles, which has social downsides.",
        "q27": "Not addressed in the app. Worth adding a note about responsible AI consumption.",
        "q29": "Improve the safety filter — add explicit escalation language for high-risk situations.",
        "q30": "Solid prototype. The ethical awareness shown in the design is commendable.",
    },
    {
        "name": "Om Singh",
        "date": "21st April 2026",
        "familiarity": "Intermediate",
        "duration": "10 minutes",
        "q1": 4, "q2": 4, "q5": 4, "q6": 3, "q8": 2, "q11": 3, "q16": 3, "q22": 2, "q28": 3,
        "q3": "The chatbot sometimes gave repetitive suggestions when I rephrased the same concern.",
        "q4": "A resource library (articles, breathing exercises) would complement the chat well.",
        "q7": "Responses were mostly relevant but occasionally generic.",
        "q9": "Minimal control. Could not go back or edit previous inputs.",
        "q10": "Only conversational text.",
        "q12": "Short note at start. No detailed policy available.",
        "q13": "No unnecessary permissions.",
        "q14": "No clear bias observed within my use.",
        "q15": "English-only currently limits accessibility for many students.",
        "q17": "Hard to distinguish AI-generated from rule-based responses.",
        "q18": "No clear feedback or error reporting channel.",
        "q19": "Limitations mentioned briefly at start — could be reinforced mid-conversation.",
        "q20": "Simple pop-up before session started.",
        "q21": "Not explicitly, but features seem passive enough to ignore.",
        "q23": "Someone in crisis might not be redirected quickly enough if the safety filter is not sensitive.",
        "q24": "No harmful outputs encountered.",
        "q25": "Students with severe anxiety or depression could develop unhealthy reliance.",
        "q26": "Might reduce visits to the counseling center, which could be net negative for serious cases.",
        "q27": "Not considered in the current version. Worth noting in documentation.",
        "q29": "Make crisis referral more prominent and immediate.",
        "q30": "Good concept and execution for a student project. Needs more robustness for real-world use.",
    },
    {
        "name": "Abhinav Gupta",
        "date": "21st April 2026",
        "familiarity": "Basic",
        "duration": "15 minutes",
        "q1": 5, "q2": 5, "q5": 4, "q6": 4, "q8": 2, "q11": 4, "q16": 4, "q22": 3, "q28": 4,
        "q3": "Nothing was confusing. Very easy to use.",
        "q4": "Maybe a daily check-in reminder notification.",
        "q7": "Responses felt natural and relevant to what I typed.",
        "q9": "I could restart the conversation but not undo a specific reply.",
        "q10": "Nothing personal — just typed how I was feeling.",
        "q12": "There was a short message saying chats are not stored. That was enough for me.",
        "q13": "No permissions asked.",
        "q14": "Did not notice any bias.",
        "q15": "Seems friendly enough for most people, but language might be a barrier.",
        "q17": "I assumed it was all AI. No indication either way.",
        "q18": "No way to flag a bad response that I could find.",
        "q19": "Yes, it said upfront it is not a real therapist.",
        "q20": "I clicked OK on a consent pop-up at the beginning.",
        "q21": "I think so, nothing felt forced.",
        "q23": "Someone very vulnerable might take the chatbot's advice too seriously.",
        "q24": "No.",
        "q25": "Younger users or those in acute distress might be at higher risk.",
        "q26": "Probably not in a major way at this stage.",
        "q27": "Did not think about it before. Worth considering.",
        "q29": "Make it clearer when to seek a real counselor.",
        "q30": "Really liked the calm tone. Felt non-judgmental.",
    },
    {
        "name": "Maulik Desai",
        "date": "18th April 2026",
        "familiarity": "Advanced",
        "duration": "20 minutes",
        "q1": 4, "q2": 4, "q5": 3, "q6": 3, "q8": 2, "q11": 3, "q16": 3, "q22": 2, "q28": 3,
        "q3": "The AI sometimes hallucinated slightly — giving advice that did not match the emotional context.",
        "q4": "Explainability layer — tell users why the AI is recommending a particular coping strategy.",
        "q7": "Yes, once it gave a breathing exercise tip when the user expressed academic stress unrelated to anxiety.",
        "q9": "Very limited. No way to adjust AI behavior or set preferences.",
        "q10": "Text messages only. However, emotional content shared could be sensitive.",
        "q12": "Minimal disclosure. For a mental health app, a full DPDP-compliant policy is needed.",
        "q13": "No unnecessary permissions, but data handling policy is underspecified.",
        "q14": "Responses occasionally reflect a Western-centric view of coping (e.g., therapy framing).",
        "q15": "Equity gaps likely for users from non-English backgrounds or low digital literacy.",
        "q17": "Not at all transparent about the decision-making process.",
        "q18": "No mechanism. Major gap for a mental health application.",
        "q19": "Disclaimer is minimal. More prominent and repeated warnings needed.",
        "q20": "One-time pop-up. Not truly informed consent for a sensitive health context.",
        "q21": "Not explicitly offered.",
        "q23": "Misinformation risk: an incorrect coping strategy recommendation could worsen someone's condition.",
        "q24": "One slightly dismissive response to a serious-sounding input. Needs better calibration.",
        "q25": "Students with clinical-level depression or suicidal ideation are the most at-risk.",
        "q26": "Could reduce demand for human counselors over time, which is a societal concern.",
        "q27": "LLM inference has a real carbon footprint. Not addressed anywhere.",
        "q29": "Implement a proper escalation protocol with real-time counselor referral.",
        "q30": "Technically sound for a prototype. Ethical depth needs significant improvement for production.",
    },
    {
        "name": "Amey Chhaya",
        "date": "21st April 2026",
        "familiarity": "Intermediate",
        "duration": "10 minutes",
        "q1": 5, "q2": 4, "q5": 4, "q6": 4, "q8": 3, "q11": 4, "q16": 4, "q22": 3, "q28": 4,
        "q3": "Nothing broken. Interface was clean and easy.",
        "q4": "Option to rate the chatbot's response in-app for feedback.",
        "q7": "Mostly accurate. One response felt a bit templated.",
        "q9": "Could stop the chat anytime, but no fine-grained control.",
        "q10": "Only text of my messages.",
        "q12": "Short consent note before the session. Simple and clear.",
        "q13": "No issues with permissions.",
        "q14": "No bias noticed.",
        "q15": "Probably yes for most students, though language and cultural context could vary.",
        "q17": "Not explicitly stated, but the responses felt AI-generated throughout.",
        "q18": "No reporting feature visible.",
        "q19": "Yes — the 'not a therapist' note at the start is important and present.",
        "q20": "Pop-up before chat with a brief explanation.",
        "q21": "Seemed optional throughout.",
        "q23": "Students might avoid seeking real help if they feel the chatbot is enough.",
        "q24": "No offensive or harmful outputs.",
        "q25": "Students already in a mental health crisis are most vulnerable.",
        "q26": "Unlikely to replace counselors fully, but could reduce first-contact visits.",
        "q27": "Minor concern at prototype scale, but worth noting for future.",
        "q29": "Add a visible 'Talk to a counselor' button throughout the chat.",
        "q30": "Liked the name — 'sahara' feels warm and approachable.",
    },
    {
        "name": "Avichal Sinha",
        "date": "18th April 2026",
        "familiarity": "Intermediate",
        "duration": "15 minutes",
        "q1": 4, "q2": 4, "q5": 4, "q6": 4, "q8": 3, "q11": 4, "q16": 4, "q22": 3, "q28": 4,
        "q3": "Slightly slow response at times. No major functional issues.",
        "q4": "Mood trend visualization over multiple sessions would be great.",
        "q7": "Responses were mostly on point. One generic reply but not harmful.",
        "q9": "Could end the session or start a new one. No finer controls.",
        "q10": "Only text chat messages.",
        "q12": "Brief disclaimer before session. Sufficient for a prototype.",
        "q13": "No unnecessary permissions.",
        "q14": "No biased outputs in my session.",
        "q15": "Should be fairly equitable for IIT student demographic. Broader use needs more testing.",
        "q17": "All responses appeared AI-driven. No distinction made.",
        "q18": "No error reporting visible.",
        "q19": "Yes, clearly stated it is not a substitute for professional help.",
        "q20": "Consent pop-up at start of session.",
        "q21": "No active opt-out feature, but prompts were not intrusive.",
        "q23": "Someone relying on it instead of calling a helpline in a real crisis is the biggest risk.",
        "q24": "No.",
        "q25": "Students going through acute mental health episodes need human support, not AI.",
        "q26": "Possible reduction in counselor visits but unlikely to eliminate the need.",
        "q27": "Not addressed in the app, but relevant for any AI system.",
        "q29": "Emergency escalation should be faster and more prominent.",
        "q30": "Well-made for a course project. Ethical framework is clearly thought through.",
    },
    {
        "name": "Ashish Aditya",
        "date": "21st April 2026",
        "familiarity": "Basic",
        "duration": "10 minutes",
        "q1": 5, "q2": 5, "q5": 5, "q6": 4, "q8": 2, "q11": 5, "q16": 4, "q22": 3, "q28": 4,
        "q3": "Nothing confusing. Very simple to use.",
        "q4": "Maybe some calming music or breathing exercises built in.",
        "q7": "The chatbot seemed to understand my mood correctly.",
        "q9": "I could stop chatting whenever I wanted.",
        "q10": "Just what I typed in the chat.",
        "q12": "There was a message at the start. I read it and it seemed fine.",
        "q13": "Nothing unnecessary was asked.",
        "q14": "Did not notice anything like that.",
        "q15": "Worked well for me. Cannot speak for everyone.",
        "q17": "Did not think about it — assumed everything was AI.",
        "q18": "I did not see any option to report a bad response.",
        "q19": "Yes — it told me it was not a real doctor.",
        "q20": "I clicked agree before starting.",
        "q21": "Nothing felt forced on me.",
        "q23": "Someone who is very sad might not realize they need real help if the app makes them feel okay.",
        "q24": "No.",
        "q25": "People going through very hard times might be affected more.",
        "q26": "Probably not much, but could reduce some load on counselors.",
        "q27": "Did not think about it before, but it is a fair point.",
        "q29": "Should always remind users to talk to a real person if things are serious.",
        "q30": "Really liked it. Felt like the app genuinely cared.",
    },
    {
        "name": "Tejas Gupta",
        "date": "21st April 2026",
        "familiarity": "Intermediate",
        "duration": "20 minutes",
        "q1": 4, "q2": 3, "q5": 4, "q6": 3, "q8": 2, "q11": 3, "q16": 3, "q22": 2, "q28": 3,
        "q3": "UI felt a bit plain. Some design polish would help trust and engagement.",
        "q4": "Session history and the ability to revisit past conversations.",
        "q7": "A few responses did not fully address what I was expressing — felt like the AI moved on too fast.",
        "q9": "Could not control depth or style of responses.",
        "q10": "Only conversational messages. No other data.",
        "q12": "Brief note at start but no detailed policy. Could be stronger.",
        "q13": "No unnecessary requests.",
        "q14": "Nothing blatant, but responses could feel disconnected from Indian student experiences.",
        "q15": "May not feel culturally relevant to all users. Localization could help.",
        "q17": "No clear distinction made between scripted and AI-generated content.",
        "q18": "No feedback mechanism in-app.",
        "q19": "Disclaimer is there but easy to miss. Should appear periodically.",
        "q20": "Pop-up consent at the beginning.",
        "q21": "Partially — no explicit controls but nothing felt mandatory.",
        "q23": "Incorrect emotional assessment by AI could make someone feel misunderstood or dismissed.",
        "q24": "Not harmful, but one response felt dismissive of a serious concern.",
        "q25": "First-time mental health help-seekers and students from non-urban backgrounds.",
        "q26": "Could partially reduce demand for peer support networks.",
        "q27": "Not considered. Would be good to add a note about this.",
        "q29": "Improve the conversational depth — the AI should follow up more carefully.",
        "q30": "Conceptually strong. Execution needs more refinement before real-world deployment.",
    },
    {
        "name": "Ankit",
        "date": "18th April 2026",
        "familiarity": "Basic",
        "duration": "10 minutes",
        "q1": 4, "q2": 4, "q5": 4, "q6": 4, "q8": 2, "q11": 4, "q16": 3, "q22": 3, "q28": 4,
        "q3": "No broken features. Interface was simple enough.",
        "q4": "Add Hindi language support for wider reach.",
        "q7": "Responses seemed relevant to what I typed.",
        "q9": "Could restart or stop. That was enough for basic use.",
        "q10": "Just the messages I sent.",
        "q12": "There was a short note before I started. Seemed okay.",
        "q13": "Nothing extra was asked.",
        "q14": "Did not notice any.",
        "q15": "English-only is a limitation. Hindi support would help a lot of students.",
        "q17": "I assumed it was AI throughout.",
        "q18": "No option to report anything.",
        "q19": "Yes, told me it is not a replacement for a doctor.",
        "q20": "I agreed to a pop-up before the chat began.",
        "q21": "Nothing felt compulsory.",
        "q23": "If the chatbot misses that someone is in crisis, that could be dangerous.",
        "q24": "No.",
        "q25": "Students who do not speak English well might get worse support.",
        "q26": "Might reduce some visits to the counselor, but real help is still needed.",
        "q27": "Not something I thought about before.",
        "q29": "Add Hindi or regional language support.",
        "q30": "Simple and calming. Good work by the team.",
    },
]

for idx, r in enumerate(respondents):
    if idx > 0:
        doc.add_page_break()

    # Header
    add_heading(doc, "Customer Feedback Questionnaire", level=1)
    p = doc.add_paragraph("AI-Based Application — Ethical Issues with AI Course")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Section A
    add_heading(doc, "Section A: Respondent Information", level=2)
    add_field(doc, "Respondent Name", r["name"])
    add_field(doc, "Date of Feedback", r["date"])
    add_field(doc, "Application Name", "Sahara")
    add_field(doc, "Team / Student Names", "Ishaan Pandey, Om Singh, Deekshant Singh Rajawat")
    add_field(doc, "How did you access the application?", "Web link")
    add_field(doc, "Duration of usage before giving feedback", r["duration"])
    add_field(doc, "Familiarity with AI-based applications", r["familiarity"])

    doc.add_paragraph()

    # Section B
    add_heading(doc, "Section B: General Usability & Functionality", level=2)
    add_rating_question(doc, 1, "How easy was it to understand the purpose of the application on first use?",
        r["q1"], ["Very Difficult", "Difficult", "Neutral", "Easy", "Very Easy"])
    add_rating_question(doc, 2, "How would you rate the overall user interface and navigation?",
        r["q2"], ["Very Poor", "Poor", "Acceptable", "Good", "Excellent"])
    add_text_question(doc, 3, "Were there any features that were confusing or did not work as expected?", r["q3"])
    add_text_question(doc, 4, "What features or functionalities would you like to see added or improved?", r["q4"])
    add_rating_question(doc, 5, "How likely are you to use this application regularly if it were fully developed?",
        r["q5"], ["Very Unlikely", "Unlikely", "Neutral", "Likely", "Very Likely"])

    doc.add_paragraph()

    # Section C
    add_heading(doc, "Section C: AI-Specific Feedback", level=2)
    add_rating_question(doc, 6, "How accurate or relevant were the AI-generated outputs/recommendations you received?",
        r["q6"], ["Very Inaccurate", "Inaccurate", "Mixed", "Accurate", "Very Accurate"])
    add_text_question(doc, 7, "Did the AI produce any outputs that were incorrect, misleading, or unexpected?", r["q7"])
    add_rating_question(doc, 8, "How well did the application explain how the AI arrived at its outputs or decisions?",
        r["q8"], ["No Explanation", "Poor", "Partial", "Good", "Fully Transparent"])
    add_text_question(doc, 9, "Did you feel you had adequate control over the AI's actions within the application?", r["q9"])

    doc.add_paragraph()

    # Section D
    add_heading(doc, "Section D: Ethical Issues", level=2)

    add_heading(doc, "D1: Privacy & Data Handling", level=2)
    add_text_question(doc, 10, "What personal data, if any, did the application ask you to provide?", r["q10"])
    add_rating_question(doc, 11, "How comfortable were you with the amount of personal data the application collected?",
        r["q11"], ["Very Uncomfortable", "Uncomfortable", "Neutral", "Comfortable", "Very Comfortable"])
    add_text_question(doc, 12, "Was there a clear privacy policy or data usage disclosure? What did it state?", r["q12"])
    add_text_question(doc, 13, "Did the application request any permissions or data that you felt were unnecessary?", r["q13"])

    add_heading(doc, "D2: Bias & Fairness", level=2)
    add_text_question(doc, 14, "Did you observe any outputs that appeared biased based on any characteristic?", r["q14"])
    add_text_question(doc, 15, "Do you think the application would produce equitable results for users from diverse backgrounds?", r["q15"])
    add_rating_question(doc, 16, "How confident are you that the application treats all users fairly?",
        r["q16"], ["Not at All", "Slightly", "Moderately", "Very", "Completely"])

    add_heading(doc, "D3: Transparency & Accountability", level=2)
    add_text_question(doc, 17, "Was it clear to you when the AI was making decisions versus a human-defined rule?", r["q17"])
    add_text_question(doc, 18, "If the AI made an error, was there a clear mechanism to report it or seek correction?", r["q18"])
    add_text_question(doc, 19, "Do you think the application adequately disclosed its limitations?", r["q19"])

    add_heading(doc, "D4: Consent & User Autonomy", level=2)
    add_text_question(doc, 20, "Did the application obtain your informed consent before collecting data or performing AI-driven actions?", r["q20"])
    add_text_question(doc, 21, "Could you opt out of specific AI features while still using the application?", r["q21"])
    add_rating_question(doc, 22, "How much control did you feel you had over your own data within the application?",
        r["q22"], ["No Control", "Little", "Some", "Good Control", "Full Control"])

    add_heading(doc, "D5: Safety & Potential Harms", level=2)
    add_text_question(doc, 23, "Can you identify any scenario where this application could cause harm to a user or a third party?", r["q23"])
    add_text_question(doc, 24, "Did you encounter any outputs that could be considered harmful, offensive, or inappropriate?", r["q24"])
    add_text_question(doc, 25, "Are there any groups of people who might be disproportionately harmed by this application?", r["q25"])

    add_heading(doc, "D6: Broader Societal & Environmental Impact", level=2)
    add_text_question(doc, 26, "Does this application have the potential to displace human jobs or reduce human decision-making?", r["q26"])
    add_text_question(doc, 27, "Are there any concerns about the environmental impact of this application?", r["q27"])

    doc.add_paragraph()

    # Section E
    add_heading(doc, "Section E: Overall Assessment", level=2)
    add_rating_question(doc, 28, "Overall, how ethically responsible do you consider this application to be?",
        r["q28"], ["Very Irresponsible", "Irresponsible", "Neutral", "Responsible", "Very Responsible"])
    add_text_question(doc, 29, "What is the single most important ethical improvement the developers should prioritize?", r["q29"])
    add_text_question(doc, 30, "Please share any additional feedback, concerns, or suggestions not covered above.", r["q30"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Declaration: ").bold = True
    p.add_run("I confirm that the feedback provided above is my honest assessment based on my experience with the application.")
    p.runs[0].font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.add_run(f"Respondent Signature: {r['name']}          Date: {r['date']}")
    p2.runs[0].font.size = Pt(10)
    p2.runs[0].italic = True

doc.save("EAI_Course_Project_Feedback_Filled.docx")
print("Done")
